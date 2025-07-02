"""
Document Processing Service for USYD Web Crawler and RAG

Handles upload, validation, and text extraction from various document formats:
- PDF files (using PyPDF2 and pdfplumber)
- Microsoft Word documents (using python-docx)  
- Markdown files (using python-markdown)

All documents are stored securely in Azure Blob Storage and processed asynchronously.
Documents use the same storage/content-passing mechanism as scraped web content.
"""

import os
import tempfile
import logging
import json
import uuid
from typing import Dict, List, Tuple, Optional
import re
from datetime import datetime

# Database
import psycopg2
from psycopg2.extras import RealDictCursor

# Azure Storage
from azure.storage.blob import BlobServiceClient
from azure.core.exceptions import AzureError

# Document processing libraries
import PyPDF2
import pdfplumber
import docx
import markdown

# File type detection
import filetype
import chardet

logger = logging.getLogger(__name__)

# Global singleton instance
_document_processor_instance = None


def get_document_processor():
    """Get singleton instance of DocumentProcessingService with
    lazy initialization"""
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessingService()
    return _document_processor_instance


# Initialize the document processor singleton
try:
    document_processor = get_document_processor()
    if document_processor.storage_enabled:
        logger.info("Document processor initialized successfully with Azure Storage")
    else:
        logger.warning("Document processor initialized but Azure Storage is not available")
        document_processor = None
except Exception as e:
    logger.error(f"Failed to initialize document processor: {e}")
    document_processor = None

class DocumentProcessingService:
    """Service for handling document upload, validation, and text extraction"""
    
    def __init__(self):
        """Initialize the document processing service with Azure Blob Storage"""
        self.db_url = os.getenv("DATABASE_URL", "postgresql://localhost:5432/usydrag")
        
        # Check if Azure Storage is configured
        storage_account_url = os.getenv("AZURE_STORAGE_ACCOUNT_URL")
        storage_key = os.getenv("AZURE_STORAGE_KEY")
        
        if storage_account_url and storage_key:
            try:
                self.blob_service_client = BlobServiceClient(
                    account_url=storage_account_url,
                    credential=storage_key
                )
                self.storage_enabled = True
                logger.info("Azure Blob Storage initialized successfully")
            except Exception as e:
                logger.warning(f"Failed to initialize Azure Blob Storage: {e}")
                self.blob_service_client = None
                self.storage_enabled = False
        else:
            logger.warning("Azure Storage not configured - document upload disabled")
            self.blob_service_client = None
            self.storage_enabled = False
            
        self.container_name = "user-documents"
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        self.allowed_extensions = ['.pdf', '.docx', '.md']
        
        # Ensure container exists only if storage is enabled
        if self.storage_enabled:
            self._ensure_container_exists()
    
    def _get_db_connection(self):
        """Get database connection"""
        try:
            conn = psycopg2.connect(self.db_url)
            return conn
        except Exception as e:
            logger.error(f"Database connection failed: {str(e)}")
            raise
    
    def _ensure_container_exists(self):
        """Ensure the Azure Blob Storage container exists"""
        try:
            container_client = self.blob_service_client.get_container_client(self.container_name)
            if not container_client.exists():
                container_client.create_container()
                logger.info(f"Created container: {self.container_name}")
        except Exception as e:
            logger.error(f"Failed to ensure container exists: {str(e)}")
    
    def validate_document(self, filename: str, file_size: int, file_content: bytes = None) -> Dict:
        """
        Validate document format, size, and content
        
        Args:
            filename: Original filename
            file_size: File size in bytes
            file_content: Optional file content for type detection
            
        Returns:
            Dict with validation result and any errors
        """
        try:
            # Check file size
            if file_size > self.max_file_size:
                return {
                    "valid": False,
                    "error": f"File size ({file_size / 1024 / 1024:.1f}MB) exceeds limit ({self.max_file_size / 1024 / 1024}MB)"
                }
            
            # Check file extension
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in self.allowed_extensions:
                return {
                    "valid": False,
                    "error": f"Unsupported file format '{file_ext}'. Allowed: {', '.join(self.allowed_extensions)}"
                }
            
            # Additional content-based validation if file content provided
            if file_content:
                # Detect actual file type
                kind = filetype.guess(file_content)
                if kind is None and file_ext != '.md':  # Markdown might not be detected
                    return {
                        "valid": False,
                        "error": "Unable to determine file type from content"
                    }
                
                # Verify file type matches extension
                if kind and file_ext == '.pdf' and kind.mime != 'application/pdf':
                    return {"valid": False, "error": "File content doesn't match PDF format"}
                elif kind and file_ext == '.docx' and 'wordprocessingml' not in kind.mime:
                    return {"valid": False, "error": "File content doesn't match Word document format"}
            
            return {"valid": True, "error": None}
            
        except Exception as e:
            logger.error(f"Validation error for {filename}: {str(e)}")
            return {"valid": False, "error": f"Validation failed: {str(e)}"}
    
    def upload_document(self, file_data: bytes, filename: str, user_id: int) -> str:
        """
        Upload document to Azure Blob Storage
        
        Args:
            file_data: File content as bytes
            filename: Original filename
            user_id: User ID for organizing files
            
        Returns:
            Blob name for the uploaded file
        """
        if not self.storage_enabled:
            raise Exception("Document upload not available - "
                            "Azure Storage not configured")
            
        try:
            # Create unique blob name
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', filename)
            blob_name = f"user_{user_id}/{timestamp}_{safe_filename}"
            
            # Upload to blob storage
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            
            blob_client.upload_blob(file_data, overwrite=True)
            logger.info(f"Uploaded document: {blob_name}")
            
            return blob_name
            
        except AzureError as e:
            logger.error(f"Azure upload error for {filename}: {str(e)}")
            raise Exception(f"Failed to upload document: {str(e)}")
        except Exception as e:
            logger.error(f"Upload error for {filename}: {str(e)}")
            raise Exception(f"Upload failed: {str(e)}")
    
    def download_document(self, blob_name: str) -> bytes:
        """Download document from Azure Blob Storage"""
        try:
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            return blob_client.download_blob().readall()
        except Exception as e:
            logger.error(f"Download error for {blob_name}: {str(e)}")
            raise Exception(f"Failed to download document: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text content from PDF file using multiple methods
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            text_content = ""
            metadata = {
                "source_type": "pdf_document",
                "extraction_method": "multiple"
            }
            
            # First try with PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Extract metadata
                    if pdf_reader.metadata:
                        metadata.update({
                            "page_count": len(pdf_reader.pages),
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                            "creator": pdf_reader.metadata.get('/Creator', '')
                        })
                    
                    # Extract text from each page
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                        except Exception as e:
                            logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                            continue
                
                # If PyPDF2 didn't extract much text, try pdfplumber
                if len(text_content.strip()) < 100:
                    logger.info("PyPDF2 extracted minimal text, trying pdfplumber")
                    with pdfplumber.open(file_path) as pdf:
                        text_content = ""
                        metadata["extraction_method"] = "pdfplumber"
                        metadata["page_count"] = len(pdf.pages)
                        
                        for page_num, page in enumerate(pdf.pages):
                            try:
                                page_text = page.extract_text()
                                if page_text:
                                    text_content += f"\n--- Page {page_num + 1} ---\n{page_text}"
                            except Exception as e:
                                logger.warning(f"pdfplumber failed on page {page_num + 1}: {str(e)}")
                                continue
                
            except Exception as e:
                logger.error(f"PDF extraction error: {str(e)}")
                raise Exception(f"Failed to extract text from PDF: {str(e)}")
            
            if not text_content.strip():
                raise Exception("No text could be extracted from PDF")
            
            # Clean up the text
            text_content = self._clean_extracted_text(text_content)
            metadata["content_length"] = len(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"PDF processing error: {str(e)}")
            raise
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text content from Word document
        
        Args:
            file_path: Path to Word document
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            doc = docx.Document(file_path)
            text_content = ""
            metadata = {
                "source_type": "word_document",
                "paragraph_count": len(doc.paragraphs)
            }
            
            # Extract core properties if available
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "created": str(core_props.created) if core_props.created else '',
                    "modified": str(core_props.modified) if core_props.modified else ''
                })
            except Exception as e:
                logger.warning(f"Could not extract document properties: {str(e)}")
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            if not text_content.strip():
                raise Exception("No text could be extracted from Word document")
            
            # Clean up the text
            text_content = self._clean_extracted_text(text_content)
            metadata["content_length"] = len(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Word document processing error: {str(e)}")
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def extract_text_from_markdown(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text content from Markdown file
        
        Args:
            file_path: Path to Markdown file
            
        Returns:
            Tuple of (text_content, metadata)
        """
        try:
            # Read file with encoding detection
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
            
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            
            if not content.strip():
                raise Exception("Markdown file is empty")
            
            # Extract frontmatter if present
            metadata = {
                "source_type": "markdown_document",
                "original_format": "markdown",
                "encoding": encoding
            }
            
            # Check for YAML frontmatter
            if content.startswith('---'):
                try:
                    end_marker = content.find('---', 3)
                    if end_marker > 0:
                        frontmatter = content[3:end_marker].strip()
                        content = content[end_marker + 3:].strip()
                        
                        # Parse basic frontmatter (title, author, etc.)
                        for line in frontmatter.split('\n'):
                            if ':' in line:
                                key, value = line.split(':', 1)
                                metadata[key.strip()] = value.strip()
                except Exception as e:
                    logger.warning(f"Could not parse frontmatter: {str(e)}")
            
            # Convert markdown to plain text (preserve some structure)
            md = markdown.Markdown(extensions=['meta', 'tables', 'fenced_code'])
            html = md.convert(content)
            
            # Simple HTML to text conversion while preserving structure
            text_content = self._html_to_text(html)
            
            if not text_content.strip():
                raise Exception("No text could be extracted from Markdown")
            
            metadata["content_length"] = len(text_content)
            
            return text_content, metadata
            
        except Exception as e:
            logger.error(f"Markdown processing error: {str(e)}")
            raise Exception(f"Failed to extract text from Markdown: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove page break artifacts
        text = re.sub(r'\n--- Page \d+ ---\n', '\n\n', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def _html_to_text(self, html: str) -> str:
        """Convert HTML to plain text while preserving structure"""
        # Replace headers with text equivalents
        html = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\n\1\n', html, flags=re.DOTALL)
        
        # Replace paragraphs
        html = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', html, flags=re.DOTALL)
        
        # Replace list items
        html = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', html, flags=re.DOTALL)
        
        # Replace line breaks
        html = re.sub(r'<br[^>]*/?>', '\n', html)
        
        # Remove all remaining HTML tags
        text = re.sub(r'<[^>]+>', '', html)
        
        # Decode HTML entities
        text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        text = text.replace('&quot;', '"').replace('&#39;', "'")
        
        return text
    
    def chunk_document_content(self, content: str, metadata: Dict, chunk_size: int = 1000) -> List[Dict]:
        """
        Split document content into chunks suitable for embedding
        
        Args:
            content: Text content to chunk
            metadata: Document metadata
            chunk_size: Target chunk size in words
            
        Returns:
            List of chunk dictionaries
        """
        chunks = []
        words = content.split()
        
        if not words:
            return chunks
        
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunk_metadata = metadata.copy()
            chunk_metadata.update({
                "chunk_index": i // chunk_size,
                "chunk_size": len(chunk_words),
                "total_chunks": (len(words) + chunk_size - 1) // chunk_size,
                "chunk_start_word": i,
                "chunk_end_word": min(i + chunk_size, len(words))
            })
            
            chunks.append({
                "content": chunk_text,
                "metadata": chunk_metadata,
                "source_url": f"document://{metadata.get('filename', 'unknown')}",
                "title": f"{metadata.get('title', metadata.get('filename', 'Document'))} - Part {chunk_metadata['chunk_index'] + 1}"
            })
        
        return chunks
    
    def process_uploaded_documents(self, blob_names: List[str], user_id: int) -> List[Dict]:
        """
        Process uploaded documents synchronously
        
        Args:
            blob_names: List of blob names to process
            user_id: User ID
            
        Returns:
            List of processed document data
        """
        processed_documents = []
        
        for blob_name in blob_names:
            try:
                logger.info(f"Processing document: {blob_name}")
                
                # Download document from blob storage
                file_data = self.download_document(blob_name)
                
                # Create temporary file
                filename = os.path.basename(blob_name)
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                    temp_file.write(file_data)
                    temp_path = temp_file.name
                
                try:
                    # Extract text based on file type
                    file_ext = os.path.splitext(filename)[1].lower()
                    
                    if file_ext == '.pdf':
                        text_content, metadata = self.extract_text_from_pdf(temp_path)
                    elif file_ext == '.docx':
                        text_content, metadata = self.extract_text_from_docx(temp_path)
                    elif file_ext == '.md':
                        text_content, metadata = self.extract_text_from_markdown(temp_path)
                    else:
                        logger.warning(f"Unsupported file type: {file_ext}")
                        continue
                    
                    # Add document info to metadata
                    metadata.update({
                        "filename": filename,
                        "blob_name": blob_name,
                        "user_id": user_id,
                        "processed_at": datetime.utcnow().isoformat()
                    })
                    
                    processed_documents.append({
                        "content": text_content,
                        "metadata": metadata,
                        "source_url": f"document://{filename}",
                        "title": metadata.get("title", filename)
                    })
                    
                    logger.info(f"Successfully processed {filename}: {len(text_content)} characters extracted")
                    
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(temp_path)
                    except Exception as e:
                        logger.warning(f"Could not delete temp file {temp_path}: {str(e)}")
                
            except Exception as e:
                logger.error(f"Error processing document {blob_name}: {str(e)}")
                continue
        
        return processed_documents
    
    def delete_document(self, blob_name: str) -> bool:
        """Delete document from Azure Blob Storage"""
        try:
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            blob_client.delete_blob(delete_snapshots="include")
            logger.info(f"Deleted document: {blob_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to delete document {blob_name}: {str(e)}")
            return False
    
    def save_document_metadata(self, user_id: int, vector_db_id: int, filename: str, 
                              original_filename: str, file_size: int, file_type: str, 
                              blob_name: str, metadata: Dict) -> int:
        """Save document metadata to database"""
        try:
            conn = self._get_db_connection()
            cursor = conn.cursor()
            
            # Create uploaded_documents table if it doesn't exist
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS uploaded_documents (
                    id SERIAL PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    vector_db_id INTEGER,
                    filename VARCHAR(255) NOT NULL,
                    original_filename VARCHAR(255) NOT NULL,
                    file_size INTEGER,
                    file_type VARCHAR(100),
                    blob_name VARCHAR(255) NOT NULL,
                    metadata JSONB,
                    processing_status VARCHAR(20) DEFAULT 'pending',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
            """)
            
            cursor.execute("""
                INSERT INTO uploaded_documents 
                (user_id, vector_db_id, filename, original_filename, file_size, 
                 file_type, blob_name, metadata, processing_status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id;
            """, (
                user_id, vector_db_id, filename, original_filename, 
                file_size, file_type, blob_name, json.dumps(metadata), 'processed'
            ))
            
            doc_id = cursor.fetchone()[0]
            conn.commit()
            cursor.close()
            conn.close()
            
            logger.info(f"Saved document metadata: {doc_id}")
            return doc_id
                
        except Exception as e:
            logger.error(f"Failed to save document metadata: {str(e)}")
            raise

    def save_processed_documents_to_storage(self, user_id: int, documents: List[Dict], file_metadata: Dict) -> str:
        """Save processed documents to storage using the same pattern as scraper and return document job ID"""
        try:
            doc_job_id = str(uuid.uuid4())
            
            # Create data directory for this document job (same pattern as scraper)
            data_dir = f"data/raw/{doc_job_id}"
            os.makedirs(data_dir, exist_ok=True)
            
            # Prepare document data in the same format as scraped data
            document_data = {
                "success": True,
                "source_type": "documents",
                "document_count": len(documents),
                "file_metadata": file_metadata,
                "results": documents,
                "processed_at": datetime.now().isoformat()
            }
            
            # Save document data to JSON file (same pattern as scraper)
            with open(f"{data_dir}/scraped_data.json", 'w', encoding='utf-8') as f:
                json.dump(document_data, f, indent=2, ensure_ascii=False)
            
            # Save metadata to database
            conn = self._get_db_connection()
            cursor = conn.cursor()
            
            # Create document_jobs table if it doesn't exist
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS document_jobs (
                    id VARCHAR(36) PRIMARY KEY,
                    user_id INTEGER REFERENCES users(id) ON DELETE CASCADE,
                    filename VARCHAR(255) NOT NULL,
                    file_size INTEGER,
                    file_type VARCHAR(100),
                    azure_blob_url VARCHAR(2048),
                    chunk_count INTEGER DEFAULT 0,
                    status VARCHAR(20) DEFAULT 'pending',
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    metadata JSONB
                );
            """)
            
            # Insert document job record
            cursor.execute("""
                INSERT INTO document_jobs 
                (id, user_id, filename, file_size, file_type, azure_blob_url, chunk_count, metadata, status)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s);
            """, (
                doc_job_id, user_id, file_metadata['filename'], 
                file_metadata['file_size'], file_metadata['file_type'],
                file_metadata.get('azure_blob_url', ''), len(documents), 
                json.dumps(file_metadata), 'completed'
            ))
            
            conn.commit()
            cursor.close()
            conn.close()
            
            logger.info(f"Saved processed documents for job {doc_job_id} to storage")
            return doc_job_id
            
        except Exception as e:
            logger.error(f"Failed to save processed documents: {str(e)}")
            raise

    def get_user_document_jobs(self, user_id: int) -> List[Dict]:
        """Get all document jobs for a user"""
        try:
            conn = self._get_db_connection()
            cursor = conn.cursor(cursor_factory=RealDictCursor)
            
            cursor.execute("""
                SELECT * FROM document_jobs 
                WHERE user_id = %s 
                ORDER BY created_at DESC;
            """, (user_id,))
            
            jobs = cursor.fetchall()
            
            cursor.close()
            conn.close()
            
            return [dict(job) for job in jobs]
            
        except Exception as e:
            logger.error(f"Failed to get user document jobs: {str(e)}")
            return []


# Create singleton instance (lazy initialization)
_document_processor_instance = None


def get_document_processor():
    """Get singleton instance of DocumentProcessingService with
    lazy initialization"""
    global _document_processor_instance
    if _document_processor_instance is None:
        _document_processor_instance = DocumentProcessingService()
    return _document_processor_instance


# For backward compatibility - Initialize the document processor singleton
try:
    document_processor = get_document_processor()
    if document_processor.storage_enabled:
        logger.info("Document processor initialized successfully with Azure Storage")
    else:
        logger.warning("Document processor initialized but Azure Storage is not available")
        document_processor = None
except Exception as e:
    logger.error(f"Failed to initialize document processor: {e}")
    document_processor = None
